import os
from datetime import datetime
from typing import Optional
from sqlalchemy.orm import Session
from fastapi import HTTPException

from app.models.models import Case, GeneratedDocument, DocumentType, DocumentFormat, CaseStatus
from app.config import settings

# Document types that represent a "final" filing deliverable. Generating one of
# these is treated as the signal that the user has finished working the case,
# so the case status is automatically advanced to COMPLETED. Reference/interim
# documents (e.g. a plain case summary) do not trigger this.
FINAL_DOCUMENT_TYPES = {
    DocumentType.COMPLAINT_DRAFT,
    DocumentType.FIR_DRAFT,
    DocumentType.LEGAL_NOTICE,
}

OUTPUT_DIR = "generated_documents"


def _ensure_output_dir() -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    return OUTPUT_DIR


def _get_case_context(case: Case) -> dict:
    """Extract context dict from case for template filling."""
    return {
        "title": case.title,
        "description": case.description,
        "category": case.category.value.title() if case.category else "Unknown",
        "subcategory": case.subcategory or "General",
        "summary": case.ai_summary or case.description[:300],
        "date": datetime.now().strftime("%d %B %Y"),
        "case_id": str(case.id),
        "laws": case.applicable_laws,
        "evidence": case.evidence_items,
        "recommendations": case.recommendations,
        "scores": case.scores,
    }


def generate_pdf_document(case: Case, doc_type: DocumentType) -> str:
    """Generate a PDF document and return the file path."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    out_dir = _ensure_output_dir()
    filename = f"case_{case.id}_{doc_type.value}_{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    filepath = os.path.join(out_dir, filename)
    ctx = _get_case_context(case)

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=18,
        textColor=colors.HexColor("#1a365d"),
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#2b4c7e"),
        spaceBefore=12,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#555555"),
    )

    story = []

    # Header
    story.append(Paragraph("GOVERNMENT LEGAL SERVICES", label_style))
    story.append(Paragraph("AI-Assisted Legal Document", label_style))
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1a365d")))
    story.append(Spacer(1, 0.4*cm))

    doc_titles = {
        DocumentType.COMPLAINT_DRAFT: "COMPLAINT DRAFT",
        DocumentType.FIR_DRAFT: "FIRST INFORMATION REPORT (DRAFT)",
        DocumentType.LEGAL_NOTICE: "LEGAL NOTICE",
        DocumentType.CASE_SUMMARY: "CASE SUMMARY REPORT",
        DocumentType.INVESTIGATION_REPORT: "INVESTIGATION REPORT",
    }
    story.append(Paragraph(doc_titles.get(doc_type, "LEGAL DOCUMENT"), title_style))
    story.append(Spacer(1, 0.4*cm))

    # Meta info table
    meta_data = [
        ["Case Reference No.", f"LA-{case.id:05d}"],
        ["Case Title", ctx["title"]],
        ["Legal Category", ctx["category"]],
        ["Subcategory", ctx["subcategory"]],
        ["Document Generated", ctx["date"]],
    ]
    meta_table = Table(meta_data, colWidths=[4*cm, 13*cm])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef2f7")),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1a365d")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("PADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.5*cm))

    # Case Description
    story.append(Paragraph("CASE DESCRIPTION", heading_style))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#2b4c7e")))
    story.append(Spacer(1, 0.2*cm))
    story.append(Paragraph(ctx["description"], body_style))
    story.append(Spacer(1, 0.3*cm))

    # AI Summary
    if ctx["summary"]:
        story.append(Paragraph("AI CASE SUMMARY", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#2b4c7e")))
        story.append(Spacer(1, 0.2*cm))
        story.append(Paragraph(ctx["summary"], body_style))
        story.append(Spacer(1, 0.3*cm))

    # Applicable Laws
    if ctx["laws"]:
        story.append(Paragraph("APPLICABLE LAWS & PROVISIONS", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#2b4c7e")))
        story.append(Spacer(1, 0.2*cm))
        for law in ctx["laws"]:
            story.append(Paragraph(
                f"<b>{law.act_name}, Section {law.section_number}</b> — {law.section_title}",
                body_style
            ))
            story.append(Paragraph(law.section_meaning, body_style))
            if law.punishment:
                story.append(Paragraph(f"<i>Punishment: {law.punishment}</i>", label_style))
            story.append(Spacer(1, 0.2*cm))

    # Evidence
    if ctx["evidence"]:
        story.append(Paragraph("EVIDENCE STATUS", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#2b4c7e")))
        story.append(Spacer(1, 0.2*cm))
        ev_data = [["Evidence Item", "Type", "Importance", "Status"]]
        for ev in ctx["evidence"]:
            ev_data.append([ev.name, ev.evidence_type, ev.importance.title(), ev.status.value.title()])
        ev_table = Table(ev_data, colWidths=[7*cm, 3*cm, 2.5*cm, 2.5*cm])
        ev_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a365d")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("PADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
        ]))
        story.append(ev_table)
        story.append(Spacer(1, 0.3*cm))

    # Recommendations
    if ctx["recommendations"]:
        story.append(Paragraph("RECOMMENDED ACTIONS", heading_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#2b4c7e")))
        story.append(Spacer(1, 0.2*cm))
        for i, rec in enumerate(ctx["recommendations"], 1):
            story.append(Paragraph(f"{i}. {rec.action}", body_style))

    # Disclaimer
    story.append(Spacer(1, 0.8*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Spacer(1, 0.2*cm))
    disclaimer = (
        "DISCLAIMER: This document has been generated by an AI-assisted legal tool. "
        "It is intended as a guide for legal professionals and government officials. "
        "This does not constitute formal legal advice. All recommendations should be "
        "reviewed by a qualified legal professional before taking action."
    )
    story.append(Paragraph(disclaimer, label_style))

    doc.build(story)
    return filepath


def generate_docx_document(case: Case, doc_type: DocumentType) -> str:
    """Generate a DOCX document and return the file path."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    out_dir = _ensure_output_dir()
    filename = f"case_{case.id}_{doc_type.value}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(out_dir, filename)
    ctx = _get_case_context(case)

    doc = DocxDocument()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("GOVERNMENT LEGAL SERVICES")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    doc_titles = {
        DocumentType.COMPLAINT_DRAFT: "COMPLAINT DRAFT",
        DocumentType.FIR_DRAFT: "FIRST INFORMATION REPORT (DRAFT)",
        DocumentType.LEGAL_NOTICE: "LEGAL NOTICE",
        DocumentType.CASE_SUMMARY: "CASE SUMMARY REPORT",
        DocumentType.INVESTIGATION_REPORT: "INVESTIGATION REPORT",
    }
    heading = doc.add_heading(doc_titles.get(doc_type, "LEGAL DOCUMENT"), level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta table
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    meta_rows = [
        ("Case Reference No.", f"LA-{case.id:05d}"),
        ("Case Title", ctx["title"]),
        ("Legal Category", ctx["category"]),
        ("Subcategory", ctx["subcategory"]),
        ("Date Generated", ctx["date"]),
    ]
    for i, (label, value) in enumerate(meta_rows):
        row = table.rows[i]
        label_cell = row.cells[0]
        label_cell.text = label
        for para in label_cell.paragraphs:
            for run in para.runs:
                run.bold = True
        row.cells[1].text = value

    doc.add_paragraph()

    # Description
    doc.add_heading("Case Description", level=2)
    doc.add_paragraph(ctx["description"])

    # AI Summary
    if ctx["summary"]:
        doc.add_heading("AI Case Summary", level=2)
        doc.add_paragraph(ctx["summary"])

    # Laws
    if ctx["laws"]:
        doc.add_heading("Applicable Laws & Provisions", level=2)
        for law in ctx["laws"]:
            p = doc.add_paragraph()
            run = p.add_run(f"Section {law.section_number}, {law.act_name} — {law.section_title}")
            run.bold = True
            doc.add_paragraph(law.section_meaning)
            if law.punishment:
                doc.add_paragraph(f"Punishment: {law.punishment}").italic = True

    # Evidence
    if ctx["evidence"]:
        doc.add_heading("Evidence Status", level=2)
        ev_table = doc.add_table(rows=1, cols=4)
        ev_table.style = "Table Grid"
        headers = ev_table.rows[0].cells
        for i, h in enumerate(["Evidence Item", "Type", "Importance", "Status"]):
            headers[i].text = h
            for para in headers[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        for ev in ctx["evidence"]:
            row = ev_table.add_row()
            row.cells[0].text = ev.name
            row.cells[1].text = ev.evidence_type
            row.cells[2].text = ev.importance.title()
            row.cells[3].text = ev.status.value.title()

    # Recommendations
    if ctx["recommendations"]:
        doc.add_heading("Recommended Actions", level=2)
        for i, rec in enumerate(ctx["recommendations"], 1):
            doc.add_paragraph(f"{i}. {rec.action}", style="List Number")

    # Disclaimer
    doc.add_paragraph()
    disclaimer_para = doc.add_paragraph(
        "DISCLAIMER: This document has been generated by an AI-assisted legal tool. "
        "It is intended as a guide for legal professionals and government officials. "
        "This does not constitute formal legal advice."
    )
    for run in disclaimer_para.runs:
        run.italic = True
        run.font.size = Pt(8)

    doc.save(filepath)
    return filepath


def generate_document(db: Session, case_id: int, doc_type: DocumentType, doc_format: DocumentFormat) -> GeneratedDocument:
    case = db.query(Case).filter(Case.id == case_id).first()
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    if doc_format == DocumentFormat.PDF:
        file_path = generate_pdf_document(case, doc_type)
    else:
        file_path = generate_docx_document(case, doc_type)

    filename = os.path.basename(file_path)
    gen_doc = GeneratedDocument(
        case_id=case_id,
        document_type=doc_type,
        document_format=doc_format,
        filename=filename,
        file_path=file_path,
    )
    db.add(gen_doc)

    # Auto-complete: once a final filing document exists for the case, and the
    # case isn't already completed, advance its status. This never happens for
    # a case that is still a DRAFT (i.e. hasn't been through AI analysis) —
    # only cases already IN_PROGRESS are auto-completed, so a document
    # generated against an unanalysed case doesn't silently skip that step.
    if doc_type in FINAL_DOCUMENT_TYPES and case.status == CaseStatus.IN_PROGRESS:
        case.status = CaseStatus.COMPLETED
        case.updated_at = datetime.utcnow()

    db.commit()
    db.refresh(gen_doc)
    return gen_doc
